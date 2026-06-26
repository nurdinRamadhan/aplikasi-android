#!/usr/bin/env python3
from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "audit_falak_ephemeris.md"
RAW = ROOT / "audit_falak_ephemeris_print_raw.docx"
OUTPUT = ROOT / "audit_falak_ephemeris_print.docx"

BODY_FONT = "Calibri"
CODE_FONT = "Cascadia Mono"
FALLBACK_CODE_FONT = "Courier New"
HEADING_COLOR = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)


def run_pandoc() -> None:
    cmd = [
        "pandoc",
        str(SOURCE),
        "--from=markdown+pipe_tables+fenced_code_blocks+smart",
        "--to=docx",
        "--standalone",
        "--toc",
        "--toc-depth=3",
        "--wrap=none",
        "--metadata",
        "lang=id-ID",
        "--output",
        str(RAW),
    ]
    subprocess.run(cmd, cwd=ROOT, check=True)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, right=80, bottom=80, left=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("right", right), ("bottom", bottom), ("left", left)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_paragraph_border(paragraph, left: bool = False, bottom: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if left:
        node = OxmlElement("w:left")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "8")
        node.set(qn("w:space"), "8")
        node.set(qn("w:color"), "000000")
        p_bdr.append(node)
    if bottom:
        node = OxmlElement("w:bottom")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "4")
        node.set(qn("w:color"), "000000")
        p_bdr.append(node)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Halaman ")
    run.font.name = BODY_FONT
    run.font.size = Pt(9)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_page_break_before(paragraph) -> None:
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    new_r.append(br)
    new_p.append(new_r)
    paragraph._p.addprevious(new_p)


def ensure_style(doc: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return doc.styles[name]
    except KeyError:
        return doc.styles.add_style(name, style_type)


def configure_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)

    for name, size, before, after in [
        ("Title", 22, 12, 10),
        ("Heading 1", 18, 18, 8),
        ("Heading 2", 14, 14, 6),
        ("Heading 3", 12, 10, 4),
        ("Heading 4", 11, 8, 3),
    ]:
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = HEADING_COLOR if name != "Heading 4" else BLACK
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        if name == "Title":
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    code = ensure_style(doc, "Audit Code Block")
    code.font.name = FALLBACK_CODE_FONT
    code._element.rPr.rFonts.set(qn("w:eastAsia"), FALLBACK_CODE_FONT)
    code.font.size = Pt(8)
    code.font.color.rgb = BLACK
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0


def configure_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.3)
        section.bottom_margin = Cm(2.3)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.2)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.1)

        header_p = section.header.paragraphs[0]
        header_p.text = "Audit Perhitungan Falak Ephemeris"
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_p.runs[0].font.name = BODY_FONT
        header_p.runs[0].font.size = Pt(9)
        header_p.runs[0].font.color.rgb = BLACK

        footer_p = section.footer.paragraphs[0]
        footer_p.clear()
        add_page_number(footer_p)


def style_paragraphs(doc: Document) -> None:
    page_break_before = {
        "Tahap 1: Markaz",
        "Workflow Fitur Hisab Hilal End-to-End",
        "Lampiran Source Excerpt Lengkap",
        "Kesimpulan Audit",
    }

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""

        if text in page_break_before:
            add_page_break_before(paragraph)

        if style_name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if style_name in {"Heading 1", "Heading 2"}:
                set_paragraph_border(paragraph, bottom=True)
            for run in paragraph.runs:
                run.font.name = BODY_FONT
                run.font.color.rgb = HEADING_COLOR if style_name != "Heading 4" else BLACK
            continue

        is_code = "Source Code" in style_name or "Code" in style_name
        if not is_code and text.startswith(("val ", "fun ", "private fun ", "data class ", "override ", "// Nama file:")):
            is_code = True

        if is_code:
            paragraph.style = doc.styles["Audit Code Block"]
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            set_paragraph_border(paragraph, left=True)
            for run in paragraph.runs:
                run.font.name = FALLBACK_CODE_FONT
                run.font.size = Pt(8)
                run.font.color.rgb = BLACK
            continue

        paragraph.paragraph_format.widow_control = True
        for run in paragraph.runs:
            run.font.name = BODY_FONT
            run.font.color.rgb = BLACK


def style_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                set_cell_margins(cell)
                if row_idx == 0:
                    set_cell_shading(cell, "FFFFFF")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(2)
                    for run in paragraph.runs:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(8.5)
                        run.font.color.rgb = BLACK
                        if row_idx == 0:
                            run.font.bold = True


def add_cover_note(doc: Document) -> None:
    first = doc.paragraphs[0] if doc.paragraphs else None
    if not first:
        return
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in first.runs:
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.color.rgb = HEADING_COLOR

    p = doc.paragraphs.insert if False else None
    # python-docx does not support inserting paragraphs before an arbitrary
    # paragraph cleanly without XML surgery; keep the source order intact.


def postprocess_docx() -> None:
    doc = Document(RAW)
    doc.core_properties.title = "Audit Perhitungan Falak Ephemeris"
    doc.core_properties.subject = "Dokumen audit dan edukasi fitur Hisab Hilal Ephemeris"
    doc.core_properties.author = "Alhasanah Media"

    configure_sections(doc)
    configure_styles(doc)
    style_paragraphs(doc)
    style_tables(doc)
    add_cover_note(doc)

    doc.save(OUTPUT)


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f"File tidak ditemukan: {SOURCE}")
    run_pandoc()
    postprocess_docx()
    print(OUTPUT)


if __name__ == "__main__":
    main()
